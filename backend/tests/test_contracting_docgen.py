from __future__ import annotations

import io

import pytest


def test_render_contract_docx_renders_and_uploads(monkeypatch):
    pytest.importorskip("docxtpl")
    pytest.importorskip("docx")

    from docx import Document

    from app.domain.pipeline.contracting import contracting_docgen

    # Build a minimal DOCX template with docxtpl placeholders.
    tmpl = Document()
    tmpl.add_paragraph("Company: {{ company.name }}")
    tmpl.add_paragraph("RFP: {{ rfp.title }}")
    tmpl.add_paragraph("Proposal: {{ proposal.title }}")
    buf = io.BytesIO()
    tmpl.save(buf)
    template_bytes = buf.getvalue()

    captured_put: dict = {}

    def fake_get_object_bytes(*, key: str, max_bytes: int = 0) -> bytes:
        assert key == "template-key"
        return template_bytes

    def fake_put_object_bytes(*, key: str, data: bytes, content_type: str | None = None):
        captured_put["key"] = key
        captured_put["data_len"] = len(data or b"")
        captured_put["content_type"] = content_type
        return {"ok": True}

    monkeypatch.setattr(contracting_docgen, "get_object_bytes", fake_get_object_bytes)
    monkeypatch.setattr(contracting_docgen, "put_object_bytes", fake_put_object_bytes)

    monkeypatch.setattr(contracting_docgen, "get_case_by_id", lambda _id: {"_id": _id, "proposalId": "p1", "rfpId": "r1", "companyId": "c1", "keyTerms": {}})
    monkeypatch.setattr(contracting_docgen, "get_proposal_by_id", lambda _id, include_sections=True: {"_id": _id, "title": "My Proposal", "companyId": "c1", "sections": {}})
    monkeypatch.setattr(contracting_docgen, "get_rfp_by_id", lambda _id: {"_id": _id, "title": "My RFP"})
    monkeypatch.setattr(contracting_docgen.content_repo, "get_company_by_company_id", lambda _id: {"companyId": _id, "name": "Acme"})

    monkeypatch.setattr(contracting_docgen, "get_contract_template", lambda _id: {"_id": _id, "currentVersionId": "v1"})
    monkeypatch.setattr(contracting_docgen, "get_contract_template_version", lambda tid, vid: {"templateId": tid, "versionId": vid, "s3Key": "template-key"})

    monkeypatch.setattr(
        contracting_docgen,
        "add_contract_doc_version",
        lambda **kwargs: {"_id": "contract_doc_1", "docxS3Key": kwargs.get("docx_s3_key")},
    )

    out = contracting_docgen.render_contract_docx(
        case_id="case_1",
        template_id="tpl_1",
        template_version_id=None,
        render_inputs={"extra": "x"},
        created_by_user_sub="u1",
    )

    assert out["version"]["_id"] == "contract_doc_1"
    assert captured_put.get("key") == out["docxS3Key"]
    assert captured_put.get("data_len", 0) > 0
    assert captured_put.get("content_type", "").startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def test_generate_budget_xlsx_generates_and_uploads(monkeypatch):
    pytest.importorskip("openpyxl")

    from app.domain.pipeline.contracting import contracting_docgen

    captured_put: dict = {}

    def fake_put_object_bytes(*, key: str, data: bytes, content_type: str | None = None):
        captured_put["key"] = key
        captured_put["data_len"] = len(data or b"")
        captured_put["content_type"] = content_type
        return {"ok": True}

    monkeypatch.setattr(contracting_docgen, "put_object_bytes", fake_put_object_bytes)
    monkeypatch.setattr(contracting_docgen, "get_case_by_id", lambda _id: {"_id": _id, "proposalId": "p1"})
    monkeypatch.setattr(contracting_docgen, "get_proposal_by_id", lambda _id, include_sections=False: {"_id": _id, "budgetBreakdown": {}})
    monkeypatch.setattr(
        contracting_docgen,
        "add_budget_version",
        lambda **kwargs: {"_id": "budget_1", "xlsxS3Key": kwargs.get("xlsx_s3_key")},
    )

    out = contracting_docgen.generate_budget_xlsx(
        case_id="case_1",
        budget_model={"currency": "USD", "items": [{"name": "PM", "rate": 100, "hours": 10}]},
        created_by_user_sub="u1",
    )

    assert out["version"]["_id"] == "budget_1"
    assert out["total"] == 1000.0
    assert captured_put.get("key") == out["xlsxS3Key"]
    assert captured_put.get("data_len", 0) > 0
    assert captured_put.get("content_type", "").startswith(
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

