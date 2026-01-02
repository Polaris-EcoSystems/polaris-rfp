from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

from fastapi import APIRouter, BackgroundTasks, HTTPException, Request
from fastapi.responses import StreamingResponse

from app.ai.client import AiError
from app.ai.user_context import load_user_profile_from_request, user_context_block
from app.ai.verified_calls import call_text_verified
from app.settings import settings
from app.infrastructure.storage import content_repo
from app.repositories import templates_repo
from app.repositories.ai_jobs_repo import create_job as create_ai_job
from app.repositories.ai_jobs_repo import update_job as update_ai_job
from app.repositories.contracting_repo import create_case, get_case_by_proposal_id
from app.repositories.rfp_proposals_repo import (
    create_proposal,
    delete_proposal,
    get_proposal_by_id,
    list_proposals,
    update_proposal,
    update_proposal_review,
)
from app.repositories.rfp_rfps_repo import get_rfp_by_id
from app.pipeline.proposal_generation.shared_section_formatters import (
    format_cover_letter_section,
    format_experience_section,
    format_title_section,
)
from app.repositories.outbox_repo import enqueue_event
from app.workflow import sync_for_rfp
from app.pipeline.proposal_generation.team_member_profiles import pick_team_member_bio, pick_team_member_experience
from app.pipeline.proposal_generation.templates_catalog import get_builtin_template, to_generator_template
from app.observability.logging import get_logger

router = APIRouter(tags=["proposals"])
log = get_logger("proposals")


def _now_iso() -> str:
    return datetime.utcnow().isoformat() + "Z"

def _clean_id_list(v: Any, *, max_items: int = 50, max_len: int = 120) -> list[str]:
    arr = v if isinstance(v, list) else []
    out: list[str] = []
    for x in arr:
        s = str(x or "").strip()
        if not s:
            continue
        s = s[:max_len]
        if s not in out:
            out.append(s)
        if len(out) >= max_items:
            break
    return out


def _generate_text_section(
    title: str,
    rfp: dict[str, Any],
    company: dict[str, Any] | None,
    team_member_ids: list[str] | None = None,
    reference_ids: list[str] | None = None,
    user_ctx: str | None = None,
) -> str:
    if not settings.openai_api_key:
        return f"{title}\n\n(This section will be completed in the proposal editor.)"

    team_ctx = ""
    if team_member_ids:
        try:
            team_ctx = _build_team_section(team_member_ids, rfp)
        except Exception:
            team_ctx = ""
    refs_ctx = ""
    if reference_ids:
        try:
            refs_ctx = _build_references_section(reference_ids)
        except Exception:
            refs_ctx = ""

    def _clip(s: str, n: int) -> str:
        s = str(s or "")
        return s if len(s) <= n else s[:n]

    prompt = (
        "Write a high-quality proposal section. Preserve markdown.\n\n"
        + (f"USER_CONTEXT:\n{user_ctx}\n\n" if user_ctx else "")
        + f"SECTION_TITLE: {title}\n"
        f"RFP_TITLE: {rfp.get('title') or ''}\n"
        f"CLIENT: {rfp.get('clientName') or ''}\n"
        f"PROJECT_TYPE: {rfp.get('projectType') or ''}\n"
        f"KEY_REQUIREMENTS: {', '.join(rfp.get('keyRequirements') or [])}\n\n"
        "COMPANY_CONTEXT:\n"
        f"- Name: {(company or {}).get('name') or ''}\n"
        f"- Capabilities: {', '.join((company or {}).get('coreCapabilities') or [])}\n\n"
        "TEAM_CONTEXT:\n"
        f"{_clip(team_ctx, 8000) if team_ctx else '(none)'}\n\n"
        "REFERENCES_CONTEXT:\n"
        f"{_clip(refs_ctx, 8000) if refs_ctx else '(none)'}\n\n"
        "Return ONLY the section content."
    )

    try:
        out, _meta = call_text_verified(
            purpose="proposal_sections",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1200,
            temperature=0.4,
            retries=2,
        )
        return out.strip() or ""
    except AiError:
        return f"{title}\n\n(This section will be completed in the proposal editor.)"


def _build_team_section(selected_ids: list[str], rfp: dict[str, Any]) -> str:
    members = content_repo.get_team_members_by_ids(selected_ids)
    members = [
        m for m in members if m and (m.get("isActive") is True or m.get("isActive") is None)
    ]
    project_type = rfp.get("projectType") if isinstance(rfp, dict) else None

    if not members:
        return "No team members selected."

    content = (
        "Our experienced team brings together diverse expertise and proven track record to deliver exceptional results.\n\n"
    )

    for member in members:
        bio = pick_team_member_bio(member, project_type)
        exp = pick_team_member_experience(member, project_type)
        content += (
            f"**{member.get('nameWithCredentials') or member.get('name') or ''}** - {member.get('position') or ''}\n\n"
        )
        if bio:
            content += f"{bio}\n\n"
        if exp:
            content += f"**Relevant experience:**\n\n{exp}\n\n"

    return content.strip()


def _build_references_section(selected_ids: list[str]) -> str:
    refs = content_repo.get_project_references_by_ids(selected_ids)
    refs = [
        r
        for r in refs
        if r
        and (r.get("isActive") is True or r.get("isActive") is None)
        and (r.get("isPublic") is True or r.get("isPublic") is None)
    ]

    if not refs:
        return "No references selected."

    content = (
        "Below are some of our recent project references that demonstrate our capabilities and client satisfaction:\n\n"
    )

    for reference in refs:
        content += f"**{reference.get('organizationName') or ''}**"
        if reference.get("timePeriod"):
            content += f" ({reference.get('timePeriod')})"
        content += "\n\n"

        content += f"**Contact:** {reference.get('contactName') or ''}"
        if reference.get("contactTitle"):
            content += f", {reference.get('contactTitle')}"
        if reference.get("additionalTitle"):
            content += f" - {reference.get('additionalTitle')}"
        content += f" of {reference.get('organizationName') or ''}\n\n"

        if reference.get("contactEmail"):
            content += f"**Email:** {reference.get('contactEmail')}\n\n"
        if reference.get("contactPhone"):
            content += f"**Phone:** {reference.get('contactPhone')}\n\n"

        content += f"**Scope of Work:** {reference.get('scopeOfWork') or ''}\n\n---\n\n"

    return content.strip()


def _section_content_from_title(
    section_title: str,
    rfp: dict[str, Any],
    company: dict[str, Any] | None,
    team_member_ids: list[str] | None = None,
    reference_ids: list[str] | None = None,
    user_ctx: str | None = None,
) -> Any:
    st = (section_title or "").lower().strip()

    if team_member_ids:
        if (
            "personnel" in st
            or "team" in st
            or "staff" in st
            or "key personnel" in st
            or "project team" in st
            or "human resource" in st
        ):
            return _build_team_section(team_member_ids, rfp)

    if reference_ids:
        if "reference" in st or "past performance" in st or "past project" in st:
            return _build_references_section(reference_ids)

    if st == "title" or st == "title page" or "title page" in st:
        return format_title_section(company, rfp)

    if "cover letter" in st or "introduction letter" in st or "transmittal letter" in st:
        return format_cover_letter_section(company, rfp)

    if (
        "experience" in st
        or "qualification" in st
        or "capabil" in st
        or "company profile" in st
        or "firm" in st
    ):
        return format_experience_section(company, rfp)

    return _generate_text_section(
        section_title,
        rfp,
        company,
        team_member_ids=team_member_ids,
        reference_ids=reference_ids,
        user_ctx=user_ctx,
    )


def _template_section_titles(template_id: str, rfp: dict[str, Any]) -> list[str]:
    tid = str(template_id or "").strip()
    if tid == "ai-template":
        raw_titles = rfp.get("sectionTitles")
        titles_in: list[Any] = raw_titles if isinstance(raw_titles, list) else []
        titles = [str(x).strip() for x in titles_in if str(x).strip()]
        if titles:
            return titles[:80]
        return [
            "Title",
            "Cover Letter",
            "Firm Qualifications and Experience",
            "Technical Approach",
            "Key Personnel",
            "References",
        ]

    builtin = get_builtin_template(tid)
    template = {**builtin, "isBuiltin": True} if builtin else templates_repo.get_template_by_id(tid)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    gen_template = to_generator_template(template)
    sec_defs = (gen_template or {}).get("sections") if isinstance(gen_template, dict) else []
    out: list[str] = []
    for s in sec_defs or []:
        name = str(s.get("title") or s.get("name") or "").strip()
        if not name:
            continue
        if name not in out:
            out.append(name)
        if len(out) >= 80:
            break
    return out or ["Title", "Cover Letter", "Technical Approach", "References"]


def _placeholder_sections(titles: list[str]) -> dict[str, Any]:
    now = _now_iso()
    out: dict[str, Any] = {}
    for t in titles:
        nm = str(t or "").strip() or "Section"
        out[nm] = {
            "content": f"{nm}\n\n(Generating…)",
            "type": "ai",
            "lastModified": now,
        }
    return out


def _render_pdf(proposal: dict[str, Any], company: dict[str, Any] | None) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except Exception as e:
        raise HTTPException(status_code=500, detail="PDF export dependency not installed") from e

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    _, height = letter

    x = 50
    y = height - 50

    c.setFont("Helvetica-Bold", 16)
    c.drawString(x, y, str(proposal.get("title") or "Proposal"))
    y -= 30

    c.setFont("Helvetica", 10)
    c.drawString(x, y, f"Company: {(company or {}).get('name') or ''}")
    y -= 20

    sections = proposal.get("sections") or {}
    for name, sec in sections.items():
        if y < 100:
            c.showPage()
            y = height - 50

        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, str(name))
        y -= 18

        c.setFont("Helvetica", 10)
        content = (sec or {}).get("content") if isinstance(sec, dict) else sec

        if isinstance(content, dict):
            content_str = "\n".join([f"{k}: {v}" for k, v in content.items()])
        else:
            content_str = str(content or "")

        for line in content_str.splitlines()[:2000]:
            if y < 60:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 10)
            c.drawString(x, y, line[:110])
            y -= 12

        y -= 10

    c.save()
    return buf.getvalue()


def _render_docx(proposal: dict[str, Any], company: dict[str, Any] | None) -> bytes:
    try:
        from docx import Document
    except Exception as e:
        raise HTTPException(status_code=500, detail="DOCX export dependency not installed") from e

    doc = Document()
    doc.add_heading(str(proposal.get("title") or "Proposal"), level=0)
    if company and company.get("name"):
        doc.add_paragraph(f"Company: {company.get('name')}")

    sections = proposal.get("sections") or {}
    for name, sec in sections.items():
        doc.add_heading(str(name), level=1)
        content = (sec or {}).get("content") if isinstance(sec, dict) else sec
        if isinstance(content, dict):
            for k, v in content.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            for line in str(content or "").splitlines():
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.post("/generate", status_code=201)
def generate(body: dict, background_tasks: BackgroundTasks, request: Request = None):  # type: ignore[assignment]
    rfp_id = (body or {}).get("rfpId")
    template_id = (body or {}).get("templateId")
    title = (body or {}).get("title")
    company_id = (body or {}).get("companyId")
    custom_content = (body or {}).get("customContent") or {}
    async_flag = bool((body or {}).get("async") is True)

    if not rfp_id or not template_id or not title:
        raise HTTPException(
            status_code=400, detail="Missing required fields: rfpId, templateId, title"
        )

    rfp = get_rfp_by_id(str(rfp_id))
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    effective_custom = dict(custom_content) if isinstance(custom_content, dict) else {}
    if company_id:
        effective_custom["companyId"] = company_id

    # Normalize selection IDs (either top-level or in customContent)
    team_member_ids = _clean_id_list(
        (body or {}).get("teamMemberIds") or effective_custom.get("teamMemberIds")
    )
    reference_ids = _clean_id_list(
        (body or {}).get("referenceIds") or effective_custom.get("referenceIds")
    )
    if team_member_ids:
        effective_custom["teamMemberIds"] = team_member_ids
    if reference_ids:
        effective_custom["referenceIds"] = reference_ids

    company = None
    if company_id:
        company = content_repo.get_company_by_company_id(str(company_id))
    if not company:
        comps = content_repo.list_companies(limit=1)
        company = comps[0] if comps else None

    titles = _template_section_titles(str(template_id), rfp)
    sections: dict[str, Any] = {}
    user_ctx = ""
    try:
        if request is not None:
            user_ctx = user_context_block(user_profile=load_user_profile_from_request(request))
    except Exception:
        user_ctx = ""

    if async_flag:
        sections = _placeholder_sections(titles)
    else:
        for t in titles:
            name = str(t)
            sections[name] = {
                "content": _section_content_from_title(
                    name,
                    rfp,
                    company,
                    team_member_ids=team_member_ids,
                    reference_ids=reference_ids,
                    user_ctx=user_ctx or None,
                ),
                "type": "ai",
                "lastModified": _now_iso(),
            }

    proposal = create_proposal(
        rfp_id=str(rfp_id),
        company_id=str(company_id) if company_id else None,
        template_id=str(template_id),
        title=str(title),
        sections=sections,
        custom_content=effective_custom,
        generation_status="queued" if async_flag else "complete",
        generation_error=None,
        generation_started_at=None,
        generation_completed_at=None if async_flag else _now_iso(),
        rfp_summary={
            "title": rfp.get("title"),
            "clientName": rfp.get("clientName"),
            "projectType": rfp.get("projectType"),
        },
    )

    if async_flag:
        proposal_id = str(proposal.get("_id") or "")

        def _run_async_generation() -> None:
            started = _now_iso()
            try:
                update_proposal(
                    proposal_id,
                    {
                        "generationStatus": "running",
                        "generationStartedAt": started,
                        "generationError": None,
                        "lastModifiedBy": "ai-generation",
                    },
                )
            except Exception:
                pass

            try:
                r = get_rfp_by_id(str(rfp_id)) or {}

                comp = None
                if company_id:
                    comp = content_repo.get_company_by_company_id(str(company_id))
                if not comp:
                    comps = content_repo.list_companies(limit=1)
                    comp = comps[0] if comps else None

                next_sections: dict[str, Any] = {}
                for t in titles:
                    nm = str(t or "").strip() or "Section"
                    next_sections[nm] = {
                        "content": _section_content_from_title(
                            nm,
                            r,
                            comp,
                            team_member_ids=team_member_ids,
                            reference_ids=reference_ids,
                            user_ctx=user_ctx or None,
                        ),
                        "type": "ai",
                        "lastModified": _now_iso(),
                    }

                done = _now_iso()
                update_proposal(
                    proposal_id,
                    {
                        "sections": next_sections,
                        "generationStatus": "complete",
                        "generationCompletedAt": done,
                        "generationError": None,
                        "lastModifiedBy": "ai-generation",
                    },
                )
            except Exception as e:
                done = _now_iso()
                try:
                    update_proposal(
                        proposal_id,
                        {
                            "generationStatus": "error",
                            "generationCompletedAt": done,
                            "generationError": (str(e) or "generation_failed")[:800],
                            "lastModifiedBy": "ai-generation",
                        },
                    )
                except Exception:
                    pass

        background_tasks.add_task(_run_async_generation)

    try:
        pid = str(proposal.get("_id") or "").strip()
        rid = str(proposal.get("rfpId") or str(rfp_id) or "").strip()
        ttl = str(proposal.get("title") or title or "").strip()
        enqueue_event(
            event_type="slack.proposal_created",
            payload={"proposalId": pid, "rfpId": rid, "title": ttl},
            dedupe_key=f"proposal_created:{pid}",
        )
    except Exception:
        # Best-effort only
        pass

    # Best-effort: sync workflow stage + seed tasks (proposal just created).
    try:
        rid = str(proposal.get("rfpId") or rfp_id or "").strip()
        if rid:
            pid2: str | None = str(proposal.get("_id") or "").strip() or None
            actor_sub: str | None = None
            try:
                if request is not None:
                    u = getattr(getattr(request, "state", None), "user", None)
                    actor_sub = str(getattr(u, "sub", "") or "").strip() if u else None
            except Exception:
                actor_sub = None
            sync_for_rfp(rfp_id=rid, actor_user_sub=actor_sub, proposal_id=pid2)
    except Exception:
        pass

    return proposal


@router.post("/{id}/generate-sections")
def generate_sections(id: str, request: Request):
    proposal = get_proposal_by_id(id, include_sections=True)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    if not settings.openai_api_key:
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")

    rfp = get_rfp_by_id(str(proposal.get("rfpId")))
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    company = None
    if proposal.get("companyId"):
        company = content_repo.get_company_by_company_id(str(proposal.get("companyId")))
    if not company:
        comps = content_repo.list_companies(limit=1)
        company = comps[0] if comps else None

    sections = proposal.get("sections") or {}
    next_sections: dict[str, Any] = {}
    user_ctx = user_context_block(user_profile=load_user_profile_from_request(request))

    for name in sections.keys():
        nm = str(name)
        next_sections[nm] = {
            "content": _section_content_from_title(nm, rfp, company, user_ctx=user_ctx or None),
            "type": "ai",
            "lastModified": _now_iso(),
        }

    updated = update_proposal(
        id, {"sections": next_sections, "lastModifiedBy": "ai-generation"}
    )

    return {"message": "Sections generated successfully", "sections": next_sections, "proposal": updated}


@router.post("/{id}/generate-sections/async")
def generate_sections_async(id: str, background_tasks: BackgroundTasks, request: Request):
    """
    Durable-ish async AI generation:
    - Creates a DynamoDB-backed job record (pollable).
    - Runs generation in a background task and updates both proposal and job status.

    Note: This is still executed within the FastAPI worker. For full durability across
    deployments/scale events, migrate the worker to a queue (SQS/Lambda/ECS worker).
    """
    proposal = get_proposal_by_id(id, include_sections=True)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    if not settings.openai_api_key:
        raise HTTPException(status_code=500, detail="OpenAI API key not configured")

    rfp = get_rfp_by_id(str(proposal.get("rfpId")))
    if not rfp:
        raise HTTPException(status_code=404, detail="RFP not found")

    company = None
    if proposal.get("companyId"):
        company = content_repo.get_company_by_company_id(str(proposal.get("companyId")))
    if not company:
        comps = content_repo.list_companies(limit=1)
        company = comps[0] if comps else None

    user = getattr(getattr(request, "state", None), "user", None)
    user_sub = getattr(user, "sub", None) if user else None
    user_ctx = user_context_block(user_profile=load_user_profile_from_request(request))

    job = create_ai_job(
        user_sub=str(user_sub) if user_sub else None,
        job_type="proposal_generate_sections",
        payload={"proposalId": str(id), "rfpId": str(proposal.get("rfpId") or "")},
    )
    job_id = str(job.get("jobId") or "").strip()

    started = _now_iso()
    try:
        update_proposal(
            id,
            {
                "generationStatus": "queued",
                "generationStartedAt": started,
                "generationError": None,
                "lastModifiedBy": "ai-generation",
            },
        )
    except Exception:
        pass

    def _run() -> None:
        try:
            update_ai_job(job_id=job_id, updates_obj={"status": "running", "startedAt": _now_iso()})
        except Exception:
            pass
        try:
            update_proposal(
                id,
                {
                    "generationStatus": "running",
                    "generationError": None,
                    "lastModifiedBy": "ai-generation",
                },
            )
        except Exception:
            pass

        try:
            sections = proposal.get("sections") or {}
            next_sections: dict[str, Any] = {}
            for name in sections.keys():
                nm = str(name)
                next_sections[nm] = {
                    "content": _section_content_from_title(nm, rfp, company, user_ctx=user_ctx or None),
                    "type": "ai",
                    "lastModified": _now_iso(),
                }

            done = _now_iso()
            update_proposal(
                id,
                {
                    "sections": next_sections,
                    "generationStatus": "complete",
                    "generationCompletedAt": done,
                    "generationError": None,
                    "lastModifiedBy": "ai-generation",
                },
            )
            try:
                update_ai_job(job_id=job_id, updates_obj={"status": "completed", "finishedAt": done, "result": {"proposalId": str(id)}})
            except Exception:
                pass
        except Exception as e:
            done = _now_iso()
            err = (str(e) or "generation_failed")[:800]
            try:
                update_proposal(
                    id,
                    {
                        "generationStatus": "error",
                        "generationCompletedAt": done,
                        "generationError": err,
                        "lastModifiedBy": "ai-generation",
                    },
                )
            except Exception:
                pass
            try:
                update_ai_job(job_id=job_id, updates_obj={"status": "failed", "finishedAt": done, "error": err})
            except Exception:
                pass

    background_tasks.add_task(_run)

    updated = get_proposal_by_id(id, include_sections=True) or proposal
    return {"ok": True, "job": job, "proposal": updated}


@router.get("/")
def list_all(request: Request, page: int = 1, limit: int = 20, nextToken: str | None = None):
    try:
        return list_proposals(page=page, limit=limit, next_token=nextToken)
    except Exception as e:
        rid = getattr(getattr(request, "state", None), "request_id", None)
        user = getattr(getattr(request, "state", None), "user", None)
        user_sub = getattr(user, "sub", None) if user else None
        log.exception(
            "proposal_list_failed",
            request_id=str(rid) if rid else None,
            user_sub=str(user_sub) if user_sub else None,
        )
        raise HTTPException(status_code=500, detail="Failed to fetch proposals") from e


@router.get("/{id}")
def get_one(id: str):
    proposal = get_proposal_by_id(id, include_sections=True)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    return proposal


@router.put("/{id}")
def update_one(id: str, request: Request, body: dict):
    before = get_proposal_by_id(id, include_sections=False) or {}
    before_status = str(before.get("status") or "").strip().lower()

    updated = update_proposal(id, {**(body or {}), "lastModifiedBy": "system"})
    if not updated:
        raise HTTPException(status_code=404, detail="Proposal not found")

    # If a proposal is marked as won, create/ensure a contracting case exists.
    try:
        after_status = str((updated or {}).get("status") or "").strip().lower()
        if after_status == "won" and before_status != "won":
            user = getattr(getattr(request, "state", None), "user", None)
            actor_sub = str(getattr(user, "sub", "") or "").strip() if user else ""
            case = get_case_by_proposal_id(str(id)) or None
            if not case:
                case = create_case(
                    proposal_id=str(id),
                    rfp_id=str((updated or {}).get("rfpId") or ""),
                    company_id=str((updated or {}).get("companyId") or "").strip() or None,
                    created_by_user_sub=actor_sub or None,
                )
            # Best-effort attach pointer to proposal for fast lookup.
            if case and case.get("_id") and not str((updated or {}).get("contractingCaseId") or "").strip():
                updated = (
                    update_proposal(
                        id,
                        {
                            "contractingCaseId": str(case.get("_id")),
                            "lastModifiedBy": "system",
                        },
                    )
                    or updated
                )
    except Exception:
        pass

    # Best-effort: sync workflow stage + seed tasks on status transitions.
    try:
        rid = str((updated or {}).get("rfpId") or "").strip()
        if rid:
            user = getattr(getattr(request, "state", None), "user", None)
            actor_sub2: str | None = str(getattr(user, "sub", "") or "").strip() if user else None
            sync_for_rfp(rfp_id=rid, actor_user_sub=actor_sub2, proposal_id=str(id))
    except Exception:
        pass
    return updated


@router.delete("/{id}")
def delete_one(id: str):
    try:
        delete_proposal(id)
        return {"message": "Proposal deleted successfully"}
    except Exception:
        raise HTTPException(status_code=500, detail="Failed to delete proposal")


@router.get("/{id}/export-pdf")
@router.get("/{id}/export/pdf")
def export_pdf(id: str):
    proposal = get_proposal_by_id(id, include_sections=True)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    company = None
    if proposal.get("companyId"):
        company = content_repo.get_company_by_company_id(str(proposal.get("companyId")))
    if not company:
        comps = content_repo.list_companies(limit=1)
        company = comps[0] if comps else None

    pdf_bytes = _render_pdf(proposal, company)
    filename = re.sub(r"\s+", "_", str(proposal.get("title") or "proposal")) + ".pdf"

    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""},
    )


@router.get("/{id}/export-docx")
def export_docx(id: str):
    proposal = get_proposal_by_id(id, include_sections=True)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    company = None
    if proposal.get("companyId"):
        company = content_repo.get_company_by_company_id(str(proposal.get("companyId")))
    if not company:
        comps = content_repo.list_companies(limit=1)
        company = comps[0] if comps else None

    docx_bytes = _render_docx(proposal, company)

    filename = re.sub(r"\s+", "_", str(proposal.get("title") or "proposal"))
    filename = re.sub(r"[^a-zA-Z0-9_-]", "", filename) + ".docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""},
    )


@router.put("/{id}/content-library/{sectionName}")
def update_content_library_section(id: str, sectionName: str, body: dict):
    selected_ids = (body or {}).get("selectedIds")
    sel = [str(x) for x in (selected_ids if isinstance(selected_ids, list) else [])]
    typ = str((body or {}).get("type") or "").strip().lower()

    proposal = get_proposal_by_id(id, include_sections=True)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    content: Any = ""

    if typ == "company":
        if sel:
            company = content_repo.get_company_by_company_id(sel[0])
            if not company:
                content = "Selected company not found."
            else:
                rfp = get_rfp_by_id(str(proposal.get("rfpId"))) or {}
                st = sectionName.lower()
                if st == "title":
                    content = format_title_section(company, rfp)
                elif "cover letter" in st or "introduction letter" in st or "transmittal letter" in st:
                    content = format_cover_letter_section(company, rfp)
                else:
                    content = format_experience_section(company, rfp)
        else:
            content = "No company selected."
    elif typ == "team":
        rfp = get_rfp_by_id(str(proposal.get("rfpId"))) or {}
        content = _build_team_section(sel, rfp)
    elif typ == "references":
        content = _build_references_section(sel)

    updated_sections = dict(proposal.get("sections") or {})
    existing_section = (
        updated_sections.get(sectionName)
        if isinstance(updated_sections.get(sectionName), dict)
        else {}
    )
    updated_sections[sectionName] = {
        **(existing_section or {}),
        "content": content.strip() if isinstance(content, str) else content,
        "type": "content-library",
        "lastModified": _now_iso(),
        "selectedIds": sel,
    }

    updated = update_proposal(id, {"sections": updated_sections})
    return updated


@router.put("/{id}/company")
def switch_company(id: str, body: dict):
    company_id = (body or {}).get("companyId")
    if not company_id or not isinstance(company_id, str):
        raise HTTPException(status_code=400, detail="companyId is required")

    proposal = get_proposal_by_id(id, include_sections=True)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    company = content_repo.get_company_by_company_id(company_id)
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")

    rfp = get_rfp_by_id(str(proposal.get("rfpId"))) or {}

    updated_sections = dict(proposal.get("sections") or {})

    updated_sections["Title"] = {
        **(updated_sections.get("Title") or {}),
        "content": format_title_section(company, rfp),
        "type": "content-library",
        "lastModified": _now_iso(),
        "selectedIds": [company_id],
    }

    updated_sections["Cover Letter"] = {
        **(updated_sections.get("Cover Letter") or {}),
        "content": format_cover_letter_section(company, rfp),
        "type": "content-library",
        "lastModified": _now_iso(),
        "selectedIds": [company_id],
    }

    for name in list(updated_sections.keys()):
        n = str(name).lower()
        if any(k in n for k in ("experience", "qualifications", "firm", "capabilities", "company profile")):
            updated_sections[name] = {
                **(updated_sections.get(name) or {}),
                "content": format_experience_section(company, rfp),
                "type": "content-library",
                "lastModified": _now_iso(),
                "selectedIds": [company_id],
            }
            break

    updated = update_proposal(
        id,
        {
            "companyId": company_id,
            "customContent": {**(proposal.get("customContent") or {}), "companyId": company_id},
            "sections": updated_sections,
            "lastModifiedBy": "system",
        },
    )
    return updated


@router.put("/{id}/review")
def update_review(id: str, body: dict):
    proposal = get_proposal_by_id(id, include_sections=True)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    score = (body or {}).get("score")
    notes = (body or {}).get("notes")
    rubric = (body or {}).get("rubric")
    decision = (body or {}).get("decision")

    next_score = None
    if score not in (None, ""):
        try:
            n = float(score) if isinstance(score, (int, float)) else float(str(score))
            next_score = max(0, min(100, n))
        except Exception:
            next_score = None

    next_decision = ((proposal.get("review") or {}).get("decision") or "")
    if decision is None:
        next_decision = ""
    if isinstance(decision, str):
        d = decision.strip().lower()
        if d in ("", "shortlist", "reject"):
            next_decision = d

    next_review = {
        **(proposal.get("review") or {}),
        "score": next_score,
        "decision": next_decision,
        "notes": notes if isinstance(notes, str) else (proposal.get("review") or {}).get("notes") or "",
        "rubric": rubric if isinstance(rubric, dict) else (proposal.get("review") or {}).get("rubric") or {},
        "updatedAt": _now_iso(),
    }

    updated = update_proposal_review(id, next_review)
    # Best-effort: stage may change based on review workflow; sync + seed tasks.
    try:
        rid = str((updated or {}).get("rfpId") or "").strip()
        if rid:
            sync_for_rfp(rfp_id=rid, actor_user_sub=None, proposal_id=str(id))
    except Exception:
        pass
    return updated
